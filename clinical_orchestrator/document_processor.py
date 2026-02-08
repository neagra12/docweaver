"""
Document Processor - Feature 2: Multi-Source Data Fusion
Handles document classification, routing, and specialized extraction
Optimized for Gemini free tier with dynamic rate limiting
"""
import asyncio
from typing import Dict, List, Any
import json
import PyPDF2
from docx import Document
from datetime import datetime, timedelta
from collections import deque

from config import Config, DOCUMENT_TYPES, get_recommended_rate_limiter_settings


class GeminiRateLimiter:
    """
    Rate limiter for Gemini API
    Dynamically configures based on model's rate limits
    """
    def __init__(self, max_calls=4, time_window=60):
        self.max_calls = max_calls
        self.time_window = time_window  # seconds
        self.call_times = deque(maxlen=max_calls)
        self.total_waits = 0
        self.total_wait_time = 0
        print(f"⚙️  Rate Limiter initialized: {max_calls} calls per {time_window}s")
    
    async def acquire(self):
        """Wait if necessary before making API call"""
        now = datetime.now()
        
        # If we have max_calls in the window, wait
        if len(self.call_times) >= self.max_calls:
            oldest_call = self.call_times[0]
            time_since_oldest = (now - oldest_call).total_seconds()
            
            if time_since_oldest < self.time_window:
                wait_time = self.time_window - time_since_oldest + 1  # +1 for safety
                print(f"⏳ Rate limit: waiting {wait_time:.1f}s before next API call...")
                await asyncio.sleep(wait_time)
                self.total_waits += 1
                self.total_wait_time += wait_time
        
        self.call_times.append(datetime.now())
    
    def get_stats(self):
        """Return rate limiting statistics"""
        return {
            "total_waits": self.total_waits,
            "total_wait_time": self.total_wait_time,
            "current_calls_in_window": len(self.call_times)
        }


# Global rate limiter instance - dynamically configured based on model
rate_limit_settings = get_recommended_rate_limiter_settings(Config.GEMINI_MODEL)
rate_limiter = GeminiRateLimiter(
    max_calls=rate_limit_settings['max_calls'],
    time_window=rate_limit_settings['time_window']
)


class DocumentProcessor:
    """Processes and classifies medical documents using Gemini"""
    
    def __init__(self):
        self.model = Config.initialize_gemini()
        self.processed_docs = []
        self.rate_limiter = rate_limiter
    
    async def classify_document(self, file_content: str) -> str:
        """
        Gemini API Call #1: Classify document type
        Routes document to appropriate specialized processor
        """
        # WAIT for rate limit before API call
        await self.rate_limiter.acquire()
        Config.increment_api_calls()
        
        prompt = f"""Analyze this medical document and classify it into ONE category:
- lab_report
- visit_note
- imaging
- specialist_note
- discharge_summary

Document content (first 1000 characters):
{file_content[:1000]}

Return ONLY the category name, nothing else."""
        
        try:
            response = await asyncio.to_thread(
                self.model.generate_content, prompt
            )
            
            doc_type = response.text.strip().lower()
            return doc_type if doc_type in DOCUMENT_TYPES else "visit_note"
        except Exception as e:
            print(f"❌ Error in classify_document: {str(e)}")
            import traceback
            traceback.print_exc()
            return "visit_note"  # Default fallback
    
    async def extract_lab_data(self, text_content: str) -> Dict[str, Any]:
        """
        Gemini API Call #2: Extract structured lab test data
        """
        await self.rate_limiter.acquire()
        Config.increment_api_calls()
        
        prompt = f"""Extract structured lab test data from this lab report.

Document:
{text_content}

Return a JSON object with this structure:
{{
    "test_date": "YYYY-MM-DD",
    "tests": [
        {{
            "name": "test name",
            "value": "numeric value",
            "unit": "unit",
            "reference_range": "normal range",
            "flag": "HIGH/LOW/NORMAL"
        }}
    ],
    "ordering_provider": "doctor name"
}}

Return ONLY valid JSON, no other text."""
        
        try:
            response = await asyncio.to_thread(
                self.model.generate_content, prompt
            )
            return self._parse_json_response(response.text)
        except Exception as e:
            print(f"❌ Error in extract_lab_data: {str(e)}")
            import traceback
            traceback.print_exc()
            return {"error": str(e)}
    
    async def extract_visit_note_data(self, text_content: str) -> Dict[str, Any]:
        """
        Gemini API Call #3: Extract visit note data
        """
        await self.rate_limiter.acquire()
        Config.increment_api_calls()
        
        prompt = f"""Extract structured data from this clinical visit note.

Document:
{text_content}

Return a JSON object:
{{
    "visit_date": "YYYY-MM-DD",
    "chief_complaint": "reason for visit",
    "diagnoses": ["diagnosis 1", "diagnosis 2"],
    "medications": ["medication 1", "medication 2"],
    "orders": ["order 1", "order 2"],
    "vital_signs": {{
        "bp": "120/80",
        "hr": "72",
        "temp": "98.6"
    }},
    "provider": "doctor name"
}}

Return ONLY valid JSON, no other text."""
        
        try:
            response = await asyncio.to_thread(
                self.model.generate_content, prompt
            )
            return self._parse_json_response(response.text)
        except Exception as e:
            print(f"❌ Error in extract_visit_note_data: {str(e)}")
            import traceback
            traceback.print_exc()
            return {"error": str(e)}
    
    async def extract_imaging_data(self, text_content: str) -> Dict[str, Any]:
        """
        Gemini API Call #4: Extract imaging report data
        """
        await self.rate_limiter.acquire()
        Config.increment_api_calls()
        
        prompt = f"""Extract structured data from this imaging report.

Document:
{text_content}

Return a JSON object:
{{
    "exam_date": "YYYY-MM-DD",
    "modality": "CT/MRI/X-ray/etc",
    "body_part": "anatomical location",
    "findings": ["finding 1", "finding 2"],
    "impression": "radiologist's conclusion",
    "recommendations": ["recommendation 1"]
}}

Return ONLY valid JSON, no other text."""
        
        try:
            response = await asyncio.to_thread(
                self.model.generate_content, prompt
            )
            return self._parse_json_response(response.text)
        except Exception as e:
            print(f"❌ Error in extract_imaging_data: {str(e)}")
            import traceback
            traceback.print_exc()
            return {"error": str(e)}
    
    async def extract_specialist_note_data(self, text_content: str) -> Dict[str, Any]:
        """
        Gemini API Call #5: Extract specialist consultation data
        """
        await self.rate_limiter.acquire()
        Config.increment_api_calls()
        
        prompt = f"""Extract structured data from this specialist consultation note.

Document:
{text_content}

Return a JSON object:
{{
    "consultation_date": "YYYY-MM-DD",
    "specialty": "specialty type",
    "reason_for_consult": "reason",
    "specialist_name": "doctor name",
    "assessment": "specialist's assessment",
    "recommendations": ["recommendation 1", "recommendation 2"],
    "follow_up": "follow-up plan"
}}

Return ONLY valid JSON, no other text."""
        
        try:
            response = await asyncio.to_thread(
                self.model.generate_content, prompt
            )
            return self._parse_json_response(response.text)
        except Exception as e:
            print(f"❌ Error in extract_specialist_note_data: {str(e)}")
            import traceback
            traceback.print_exc()
            return {"error": str(e)}
    
    async def process_document(self, file_name: str, file_content: str, doc_index: int = 0, total_docs: int = 1) -> Dict[str, Any]:
        """
        Main document processing pipeline
        1. Classify document type
        2. Route to specialized extractor
        3. Return structured data
        """
        print(f"\n📄 Processing document {doc_index + 1}/{total_docs}: {file_name}")
        
        # Classify document (API Call #1)
        print(f"   🔍 Classifying document type...")
        doc_type = await self.classify_document(file_content)
        print(f"   ✓ Identified as: {doc_type}")
        
        # Route to specialized extractor (API Calls #2-5)
        print(f"   📊 Extracting structured data...")
        extracted_data = None
        if doc_type == "lab_report":
            extracted_data = await self.extract_lab_data(file_content)
        elif doc_type == "visit_note":
            extracted_data = await self.extract_visit_note_data(file_content)
        elif doc_type == "imaging":
            extracted_data = await self.extract_imaging_data(file_content)
        elif doc_type == "specialist_note":
            extracted_data = await self.extract_specialist_note_data(file_content)
        elif doc_type == "discharge_summary":
            # Use visit note extractor for discharge summaries
            extracted_data = await self.extract_visit_note_data(file_content)
        else:
            # Default fallback
            extracted_data = await self.extract_visit_note_data(file_content)
        
        print(f"   ✓ Extraction complete")
        
        result = {
            "file_name": file_name,
            "document_type": doc_type,
            "extracted_data": extracted_data,
            "raw_text": file_content[:500],  # First 500 chars for reference
            "processing_timestamp": datetime.now().isoformat()
        }
        
        self.processed_docs.append(result)
        return result
    
    async def process_multiple_documents(self, documents: List[Dict[str, str]]) -> List[Dict[str, Any]]:
        """
        Process multiple documents SEQUENTIALLY (not parallel)
        Required for Gemini free tier rate limiting
        
        This is still sophisticated orchestration because:
        1. Each document requires 2 API calls (classify + extract)
        2. Intelligent routing to specialized processors
        3. Rate limiting coordination across all calls
        4. State management across document set
        """
        print(f"\n🚀 Starting sequential document processing ({len(documents)} documents)")
        print(f"⚙️  Rate limit: {self.rate_limiter.max_calls} calls per {self.rate_limiter.time_window}s")
        
        results = []
        start_time = datetime.now()
        
        for i, doc in enumerate(documents):
            result = await self.process_document(
                doc['file_name'], 
                doc['content'],
                doc_index=i,
                total_docs=len(documents)
            )
            results.append(result)
        
        end_time = datetime.now()
        processing_time = (end_time - start_time).total_seconds()
        
        # Get rate limiting stats
        rate_stats = self.rate_limiter.get_stats()
        
        print(f"\n✅ All documents processed!")
        print(f"   Total time: {processing_time:.1f}s")
        print(f"   API calls made: {Config.get_api_call_count()}")
        print(f"   Rate limit waits: {rate_stats['total_waits']}")
        print(f"   Total wait time: {rate_stats['total_wait_time']:.1f}s")
        
        return results
    
    def _parse_json_response(self, response_text: str) -> Dict[str, Any]:
        """Parse JSON from Gemini response, handling markdown code blocks"""
        try:
            json_text = response_text.strip()
            # Remove markdown code blocks if present
            if json_text.startswith("```json"):
                json_text = json_text[7:]
            if json_text.startswith("```"):
                json_text = json_text[3:]
            if json_text.endswith("```"):
                json_text = json_text[:-3]
            json_text = json_text.strip()
            
            return json.loads(json_text)
        except Exception as e:
            print(f"⚠️  JSON parsing warning: {str(e)}")
            return {"error": f"Failed to parse JSON: {str(e)}", "raw": response_text[:200]}
    
    def get_processing_summary(self) -> Dict[str, Any]:
        """Return summary of all processed documents"""
        return {
            "total_documents": len(self.processed_docs),
            "document_types": [doc["document_type"] for doc in self.processed_docs],
            "api_calls_used": Config.get_api_call_count(),
            "rate_limit_stats": self.rate_limiter.get_stats(),
            "documents": self.processed_docs
        }


# Utility function to extract text from files
def extract_text_from_file(file_path: str) -> str:
    """Extract text from PDF, DOCX, or TXT files"""
    if file_path.endswith('.pdf'):
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text()
                return text
        except Exception as e:
            return f"Error reading PDF: {str(e)}"
    
    elif file_path.endswith('.docx'):
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            return f"Error reading DOCX: {str(e)}"
    
    elif file_path.endswith('.txt'):
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            return f"Error reading TXT: {str(e)}"
    
    return ""


def extract_text_from_uploaded_file(uploaded_file) -> str:
    """Extract text from Streamlit uploaded file object"""
    try:
        file_extension = uploaded_file.name.split('.')[-1].lower()
        
        if file_extension == 'pdf':
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()
            return text
        
        elif file_extension == 'docx':
            doc = Document(uploaded_file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        
        elif file_extension == 'txt':
            return uploaded_file.read().decode('utf-8')
        
        else:
            return f"Unsupported file type: {file_extension}"
    
    except Exception as e:
        return f"Error reading file: {str(e)}"